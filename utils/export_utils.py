#  utils/export_utils.py

import os
import datetime
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def export_to_word(content: str, topic: str = "Untitled") -> str:
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_topic = topic.replace(' ', '_').replace('/', '_')
    filename = f"{safe_topic}_{timestamp}.docx"
    output_dir = "outputs"
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)

    doc = Document()
    doc.add_heading(topic, 0)
    for line in content.split("\n"):
        doc.add_paragraph(line)

    print(f"Saving Word file to: {filepath}")
    doc.save(filepath)

    # ✅ No auto-open here!
    return f"outputs/{filename}"

def export_to_pdf(content: str, topic: str = "Untitled") -> str:
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_topic = topic.replace(' ', '_').replace('/', '_')
    filename = f"{safe_topic}_{timestamp}.pdf"
    output_dir = "outputs"
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)

    c = canvas.Canvas(filepath, pagesize=letter)
    width, height = letter
    y = height - 50
    lines = content.split('\n')
    c.setFont("Helvetica", 11)

    for line in lines:
        if y < 50:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 11)
        c.drawString(50, y, line[:120])
        y -= 15

    print(f"Saving PDF file to: {filepath}")
    c.save()

    # ✅ No auto-open here!
    return f"outputs/{filename}"
